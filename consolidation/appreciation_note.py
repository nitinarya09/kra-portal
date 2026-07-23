"""
appreciation_note.py -- Auto-generate the Appreciation Note paragraphs
============================================================================
Reads key metrics from the downloaded KRA data and generates the
Appreciation Note text that appears at the end of the KRA report.

Usage:
    from appreciation_note import generate_appreciation_note
    note_text = generate_appreciation_note(data, fy, quarter)
"""


def _safe_float(val, default=0):
    """Convert to float safely."""
    try:
        return float(val)
    except (ValueError, TypeError):
        return default


def _get_first(data, ws_name, **filters):
    """Get first matching record from a worksheet."""
    records = data.get(ws_name, [])
    for r in records:
        match = all(
            str(r.get(k, "")).strip().lower() == str(v).strip().lower()
            for k, v in filters.items()
        )
        if match:
            return r
    return None


# Quarter -> months mapping
QUARTER_MONTHS = {
    "Q1": ("March", "April", "May"),
    "Q2": ("June", "July", "August"),
    "Q3": ("September", "October", "November"),
    "Q4": ("December", "January", "February"),
}


def generate_appreciation_note(data, fy, quarter):
    """
    Generate the Appreciation Note text for the KRA report.

    Args:
        data: dict from fetch_data — {worksheet_name: [records]}
        fy: Financial Year string, e.g. "2026-27"
        quarter: Quarter string, e.g. "Q1"

    Returns:
        list of str: Each string is one paragraph/bullet point of the note.
    """
    months = QUARTER_MONTHS.get(quarter, ("", "", ""))
    month1, month2, month3 = months
    fy_year = fy.split("-")[0] if "-" in fy else fy

    paragraphs = []

    # Header
    paragraphs.append(f"KRA report for the quarter ending {_quarter_ending(quarter)} - {fy_year}")
    paragraphs.append("")  # blank line

    # 1. MCA submission
    mca = data.get("Accounts_MCA", [])
    if mca:
        m = mca[0]
        m1 = str(m.get("Month1", month1))
        m2 = str(m.get("Month2", month2))
        m3 = str(m.get("Month3", month3))
        delay1 = str(m.get("Delay1", "Nil"))
        delay2 = str(m.get("Delay2", "Nil"))
        delay3 = str(m.get("Delay3", "Nil"))
        if all(d in ("0", "Nil", "nil", "", "0.0") for d in [delay1, delay2, delay3]):
            paragraphs.append(
                f"Monthly Civil Accounts ({m1}-{fy_year}, {m2}-{fy_year} and {m3}-{fy_year}) "
                f"have been submitted to the State Government within the specified time."
            )
        else:
            paragraphs.append(
                f"Monthly Civil Accounts ({m1}-{fy_year}, {m2}-{fy_year} and {m3}-{fy_year}) "
                f"have been submitted. Delays: {m1}: {delay1}, {m2}: {delay2}, {m3}: {delay3}."
            )

    # 2. MKI uploaded on time
    mki = data.get("MKI_Dates", [])
    if mki:
        paragraphs.append(
            "Monthly Key Indicators have been uploaded on the CAG's website, within 5 working days."
        )

    # 3. Reconciliation
    recon = data.get("Accounts_Reconciliation", [])
    if recon:
        exp_pct = ""
        rcpt_pct = ""
        for r in recon:
            section = str(r.get("Section", ""))
            if "Expenditure" in section or "3a" in section.lower():
                if str(r.get("Sector", "")).lower() == "treasury":
                    exp_pct = str(r.get("PctReconciled", ""))
            elif "Receipts" in section or "3b" in section.lower():
                if str(r.get("Sector", "")).lower() == "treasury":
                    rcpt_pct = str(r.get("PctReconciled", ""))
        if exp_pct or rcpt_pct:
            paragraphs.append(
                f"The DTA has reconciled the figures. "
                f"The percentage of reconciliation of Expenditure is {exp_pct}% "
                f"and Receipts is {rcpt_pct}%."
            )

    # 4. RTI disposal
    complaints = data.get("Complaints_RTI", [])
    rti_ok = True
    for r in complaints:
        if "rti" in str(r.get("Type", "")).lower():
            beyond = _safe_float(r.get("DisposedBeyond", 0))
            if beyond > 0:
                rti_ok = False
    if rti_ok:
        paragraphs.append(
            "All Cases of RTI have been disposed of within specified time period under RTI Act."
        )

    # 5. CV Validation
    cv = data.get("Accounts_Vouchers", [])
    if cv:
        r = cv[0]
        paragraphs.append(
            f"Primary Validation and Complete Validation of vouchers have been done. "
            f"Monthly Validation Reports have been issued to State Government."
        )

    # 6. FP Cases
    gpf = data.get("GPF_Summary", [])
    fp_rec = None
    rb_rec = None
    mc_rec = None
    mb_rec = None
    for r in gpf:
        item = str(r.get("KRA_Item", "")).lower()
        if "final" in item:
            fp_rec = r
        elif "residual" in item:
            rb_rec = r
        elif "missing" in item and "credit" in item:
            mc_rec = r
        elif "minus" in item:
            mb_rec = r

    if fp_rec:
        total = _safe_float(fp_rec.get("Total"))
        cleared = _safe_float(fp_rec.get("Clearance"))
        pct = _safe_float(fp_rec.get("ClearPct"))
        if total > 0:
            paragraphs.append(
                f"FP Cases: - {int(cleared)} out of {int(total)} FP Cases were cleared "
                f"within citizen charter's time frame ({pct:.2f}% clearance)."
            )

    if rb_rec:
        cleared = _safe_float(rb_rec.get("Clearance"))
        pct = _safe_float(rb_rec.get("ClearPct"))
        paragraphs.append(
            f"RB Cases: - {int(cleared)} RB Cases were cleared within citizen charter's "
            f"time frame ({pct:.2f}% clearance)."
        )

    # 7. Missing Credits and Unposted Items
    if mc_rec:
        mc_pct = _safe_float(mc_rec.get("ClearPct"))
        paragraphs.append(
            f"Missing credits and Unposted items: - During the reporting period, "
            f"{mc_pct:.2f}% of missing credits have been successfully cleared against the "
            f"quarterly target. GPF Camps are being organized constantly and that has yielded "
            f"good outcome."
        )

    # 8. Minus Balance
    if mb_rec:
        mb_cleared = _safe_float(mb_rec.get("Clearance"))
        mb_pct = _safe_float(mb_rec.get("ClearPct"))
        mb_target = _safe_float(mb_rec.get("AnnualTarget", mb_rec.get("Target", 0)))
        paragraphs.append(
            f"Minus Balance Cases: - Due to sincere efforts and constant correspondence "
            f"with DTA, DDOs and Treasury Officers, {int(mb_cleared)} live minus cases "
            f"({mb_pct:.2f}%) were cleared against the quarterly target of "
            f"{int(mb_target)}."
        )

    # 9. GPF Posting Arrears
    misc = data.get("GPF_Misc", [])
    posting_month = ""
    for r in misc:
        field = str(r.get("Field", "")).lower()
        if "posting" in field and "done" in field:
            posting_month = str(r.get("Value", ""))
    if posting_month:
        paragraphs.append(
            f"Arrear in Posting of GPF Accounts: - Posting till month "
            f"{posting_month} has been completed."
        )

    # 10. Complaint disposal
    total_cleared = 0
    total_cag = 0
    total_direct = 0
    for r in complaints:
        comp_type = str(r.get("Type", "")).lower()
        disposed = _safe_float(r.get("DisposedInTime", 0))
        if "cag" in comp_type:
            total_cag += disposed
        elif "direct" in comp_type:
            total_direct += disposed
        total_cleared += disposed

    if total_cleared > 0:
        paragraphs.append(
            f"Disposal of Complaint Cases: - {int(total_cag)} CAG Complaints and "
            f"{int(total_direct)} direct complaints (Total = {int(total_cleared)}) "
            f"are cleared within citizen charter's time frame."
        )

    return paragraphs


def _quarter_ending(quarter):
    """Map quarter code to ending month name."""
    endings = {"Q1": "June", "Q2": "September", "Q3": "December", "Q4": "March"}
    return endings.get(quarter, quarter)


def insert_appreciation_note(doc, note_paragraphs):
    """
    Append the Appreciation Note paragraphs after the last table in the document.

    Args:
        doc: python-docx Document object
        note_paragraphs: list of strings from generate_appreciation_note()
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Add a blank paragraph separator
    doc.add_paragraph("")

    # Add heading
    heading = doc.add_paragraph("Appreciation Note")
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = "Arial"

    # Add each bullet point
    for text in note_paragraphs:
        if not text:
            doc.add_paragraph("")
            continue
        para = doc.add_paragraph(f"•  {text}")
        for run in para.runs:
            run.font.size = Pt(10)
            run.font.name = "Arial"


# ---------------------------------------------------------------------------
# CLI entry point for testing
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    print("Appreciation Note Generator")
    print("This module is meant to be imported by consolidate.py")
    print("")
    print("Template paragraphs (with placeholder data):")
    print("-" * 60)

    # Demo with empty data
    sample = generate_appreciation_note({}, "2026-27", "Q1")
    for p in sample:
        print(f"  {p}")
