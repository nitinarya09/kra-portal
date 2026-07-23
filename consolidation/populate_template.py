"""
populate_template.py -- Fill all 48 tables in the KRA Word template
============================================================================
Opens the blank KRA Word template (.docx), locates each table by index,
and fills in cell values from the downloaded Google Sheets data.

Usage:
    from populate_template import populate_all_tables
    populate_all_tables(doc, data)
"""

from docx import Document
from docx.shared import Pt
import copy


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _safe(val, default=""):
    """Return val as string, or default if None/empty."""
    if val is None or str(val).strip() == "":
        return default
    return str(val).strip()


def _num(val, decimals=2):
    """Format a number to fixed decimals, or return as-is if not numeric."""
    try:
        f = float(val)
        if f == int(f) and decimals == 0:
            return str(int(f))
        return f"{f:,.{decimals}f}"
    except (ValueError, TypeError):
        return _safe(val, "")


def _pct(val):
    """Format a percentage value, appending % if purely numeric."""
    s = _safe(val, "")
    if not s:
        return ""
    try:
        f = float(s)
        return f"{f:.2f}%"
    except (ValueError, TypeError):
        return s if "%" in s else f"{s}%"


def set_cell(table, row_idx, col_idx, value):
    """
    Set a table cell's text while preserving the cell's paragraph formatting.
    Preserves font name, size, bold, alignment from the first run.
    """
    try:
        cell = table.cell(row_idx, col_idx)
    except IndexError:
        return  # Skip if cell doesn't exist

    text = _safe(value, "")

    # Preserve formatting from existing run if available
    para = cell.paragraphs[0]
    if para.runs:
        run = para.runs[0]
        font_name = run.font.name
        font_size = run.font.size
        font_bold = run.font.bold
    else:
        font_name = "Arial"
        font_size = Pt(9)
        font_bold = False

    # Clear existing text
    para.clear()

    # Add new run with preserved formatting
    new_run = para.add_run(text)
    new_run.font.name = font_name or "Arial"
    new_run.font.size = font_size or Pt(9)
    if font_bold:
        new_run.font.bold = True


def _find_row_by_text(table, col_idx, search_text, start_row=0, end_row=None):
    """Find a row index where cell text contains search_text."""
    if end_row is None:
        end_row = len(table.rows)
    for r in range(start_row, end_row):
        try:
            cell_text = table.cell(r, col_idx).text.strip()
            if search_text.lower() in cell_text.lower():
                return r
        except IndexError:
            continue
    return None


def _get_record(data, ws_name, **filters):
    """Get the first record matching optional key=value filters."""
    records = data.get(ws_name, [])
    for r in records:
        match = all(
            _safe(r.get(k, "")).lower() == _safe(v, "").lower()
            for k, v in filters.items()
        )
        if match:
            return r
    return None


def _get_records(data, ws_name, **filters):
    """Get all records matching optional key=value filters."""
    records = data.get(ws_name, [])
    result = []
    for r in records:
        match = all(
            _safe(r.get(k, "")).lower() == _safe(v, "").lower()
            for k, v in filters.items()
        )
        if match:
            result.append(r)
    return result


# ---------------------------------------------------------------------------
# Table 0: Main KRA Summary (Rows 0-23)
# Items 1-7: MCA, Exclusion, Reconciliation, RBD, CV, AC Bills, UCs
# ---------------------------------------------------------------------------
def populate_table_0(table, data):
    """Fill Table 0 — the primary summary table."""

    # --- Item 1: MCA rendition delay (Row 4) ---
    mca = data.get("Accounts_MCA", [])
    if mca:
        m = mca[0]
        # This Quarter columns start at col 2
        set_cell(table, 4, 3, _safe(m.get("Date1")))
        set_cell(table, 4, 4, _safe(m.get("Delay1"), "Nil"))
        # Item 2: Exclusion (Row 5)
        set_cell(table, 5, 3, _pct(m.get("ExclPct1")))

    # --- Item 3(a): Reconciliation of Expenditure (Rows 7-9) ---
    recon = data.get("Accounts_Reconciliation", [])
    sector_col_map = {"Treasury": 3, "PWD": 4, "Forest": 5}
    for r in recon:
        sector = _safe(r.get("Sector"))
        col = sector_col_map.get(sector)
        if col is None:
            continue
        # Check if this is expenditure or receipts based on Section code
        section = _safe(r.get("Section", ""))
        if "Receipts" in section or "3b" in section.lower():
            # Item 3(b): Receipts (Rows 10-12)
            set_cell(table, 10, col, _num(r.get("AmtDue")))
            set_cell(table, 11, col, _num(r.get("AmtReconciled")))
            set_cell(table, 12, col, _pct(r.get("PctReconciled")))
        else:
            # Item 3(a): Expenditure (Rows 7-9)
            set_cell(table, 7, col, _num(r.get("AmtDue")))
            set_cell(table, 8, col, _num(r.get("AmtReconciled")))
            set_cell(table, 9, col, _pct(r.get("PctReconciled")))

    # --- Item 4: RBD (Rows 13-17) ---
    rbd = data.get("Accounts_RBD", [])
    if rbd:
        r = rbd[0]
        set_cell(table, 14, 2, _safe(r.get("MonthReconciled")))
        set_cell(table, 14, 3, _safe(r.get("ArrearMonths"), "-"))
        set_cell(table, 15, 2, _num(r.get("RBI_Amt")))
        set_cell(table, 16, 2, _num(r.get("AG_Amt")))
        set_cell(table, 17, 2, _num(r.get("NetDiff")))

    # --- Item 5: CV Voucher Check (Rows 18-19) ---
    cv = data.get("Accounts_Vouchers", [])
    if cv:
        r = cv[0]
        set_cell(table, 19, 1, _num(r.get("TotalReceived"), 0))
        set_cell(table, 19, 2, _num(r.get("TotalDue"), 0))
        set_cell(table, 19, 3, _pct(r.get("PctChecked")))

    # --- Item 6: AC Bills (Rows 20-21) ---
    ac = data.get("Accounts_ACBills", [])
    if ac:
        r = ac[0]
        set_cell(table, 21, 2, _num(r.get("OB")))
        set_cell(table, 21, 3, _num(r.get("Addition")))
        set_cell(table, 21, 4, _num(r.get("Total")))
        set_cell(table, 21, 5, f"{_num(r.get('Clearance'))} ({_pct(r.get('ClearPct'))})")
        set_cell(table, 21, 6, _num(r.get("CB")))

    # --- Item 7: UCs (Rows 22-23) ---
    uc = data.get("Accounts_UCs", [])
    # Sum all year-wise rows for the summary
    if uc:
        total_ob = sum(float(r.get("OB", 0) or 0) for r in uc)
        total_add = sum(float(r.get("Addition", 0) or 0) for r in uc)
        total_total = sum(float(r.get("Total", 0) or 0) for r in uc)
        total_clear = sum(float(r.get("Clearance", 0) or 0) for r in uc)
        total_cb = sum(float(r.get("CB", 0) or 0) for r in uc)
        set_cell(table, 23, 1, _num(total_ob))
        set_cell(table, 23, 2, _num(total_add))
        set_cell(table, 23, 3, _num(total_total))
        set_cell(table, 23, 4, _num(total_clear))
        set_cell(table, 23, 5, _num(total_cb))


# ---------------------------------------------------------------------------
# Table 1: Suspense & Remittance Balances (Rows 0-19)
# Items 8(a) through 8(i)
# ---------------------------------------------------------------------------
def populate_table_1(table, data):
    """Fill Table 1 — Suspense and Remittance heads."""
    suspense = data.get("Suspense_Remittance", [])

    # Mapping: (Head contains, Dr_Cr) -> row index in Table 1
    row_map = {
        ("8658-101", "Dr"): 2,   # PAO Suspense Dr
        ("8658-101", "Cr"): 3,   # PAO Suspense Cr
        ("8658-102", "Dr"): 4,   # Suspense Civil Dr
        ("8658-102", "Cr"): 5,   # Suspense Civil Cr
        ("8658-109", "Dr"): 6,   # RBS HQ Dr
        ("8658-109", "Cr"): 7,   # RBS HQ Cr
        ("8658-110", "Dr"): 8,   # RBS CAO Dr
        ("8658-110", "Cr"): 9,   # RBS CAO Cr
        ("8658-111", "Dr"): 10,  # DAA Suspense Dr
        ("8658-111", "Cr"): 11,  # DAA Suspense Cr
        ("8782-102", "Dr"): 12,  # PW Remittance Dr
        ("8782-102", "Cr"): 13,  # PW Remittance Cr
        ("8782-103", "Dr"): 14,  # Forest Remittance Dr
        ("8782-103", "Cr"): 15,  # Forest Remittance Cr
        ("8793",     "Dr"): 16,  # Inter State Suspense Dr
        ("8793",     "Cr"): 17,  # Inter State Suspense Cr
    }

    for rec in suspense:
        head = _safe(rec.get("Head", ""))
        dr_cr = _safe(rec.get("Dr_Cr", rec.get("SubHead", ""))).strip()

        # Find matching row
        target_row = None
        for (head_pattern, dc), row_idx in row_map.items():
            if head_pattern in head and dc.lower() == dr_cr.lower():
                target_row = row_idx
                break

        if target_row is None:
            continue

        # Columns: 4=OB, 5=Addition, 6=Total, 7=Clearance+ClearPct, 8=CB
        set_cell(table, target_row, 4, _num(rec.get("OB")))
        set_cell(table, target_row, 5, _num(rec.get("Addition")))
        set_cell(table, target_row, 6, _num(rec.get("Total")))
        clearance_text = _num(rec.get("Clearance"))
        clear_pct = rec.get("ClearPct")
        if clear_pct:
            clearance_text += f"  {_pct(clear_pct)}"
        set_cell(table, target_row, 7, clearance_text)
        set_cell(table, target_row, 8, _num(rec.get("CB")))
        # Old balances cleared % (col 9 if present)
        old_cleared = rec.get("OldCleared")
        if old_cleared:
            set_cell(table, target_row, 9, _pct(old_cleared))


# ---------------------------------------------------------------------------
# Table 2: Items 9-15 (MKI, Report dates)
# ---------------------------------------------------------------------------
def populate_table_2(table, data):
    """Fill Table 2 — MKI hosting dates and Report section items 10-15."""

    # Row 0: MKI Dates (Item 9)
    mki = data.get("MKI_Dates", [])
    for i, rec in enumerate(mki[:3]):
        col_offset = 2 + i  # Cols 2, 3, 4
        set_cell(table, 0, col_offset, _safe(rec.get("UploadDate")))

    # Rows 1-7: Report items 10-15
    annual = data.get("AnnualAccounts", [])
    for rec in annual:
        item = _safe(rec.get("Item", ""))
        date1 = _safe(rec.get("Date1", rec.get("Description", "")))
        date2 = _safe(rec.get("Date2", ""))

        # Map item text to row
        if "NTFA" in item and "Audit" in item:
            set_cell(table, 1, 3, date1)
        elif "Statements" in item or "Appendices" in item:
            set_cell(table, 1, 3, date1)  # Appended on same row
        elif "Grants" in item:
            set_cell(table, 1, 3, date1)
        elif "GA Wing" in item:
            set_cell(table, 2, 3, date1)
        elif "spiral" in item.lower():
            set_cell(table, 3, 3, date1)
            if date2:
                set_cell(table, 3, 5, date2)
        elif "Closure" in item:
            set_cell(table, 5, 2, date1)
        elif "NTA" in item:
            set_cell(table, 6, 2, date1)
        elif "Signature" in item or "CAG" in item:
            set_cell(table, 7, 3, date1)


# ---------------------------------------------------------------------------
# Table 3: Items 16-20 (LTA, Loan monitoring, TI, Book-2)
# ---------------------------------------------------------------------------
def populate_table_3(table, data):
    """Fill Table 3 — LTA, Loan monitoring, TI inspection, Budget review."""

    # Item 16: LTA Posting Arrears (Rows 1-3)
    lta = data.get("LTA_Loans", [])
    for rec in lta:
        lta_type = _safe(rec.get("Type", ""))
        if "posting" in lta_type.lower() or "month" in lta_type.lower():
            set_cell(table, 3, 2, _safe(rec.get("PostingDone"), "Nil"))
            set_cell(table, 3, 3, _safe(rec.get("Arrear"), "Nil"))

    # Item 17: Loan monitoring (Rows 4-6)
    for rec in lta:
        if "principal" in _safe(rec.get("Type", "")).lower():
            set_cell(table, 5, 2, _num(rec.get("PrincipalAmt")))
        elif "interest" in _safe(rec.get("Type", "")).lower():
            set_cell(table, 6, 2, _num(rec.get("InterestAmt")))

    # Item 18: TI Inspection (Rows 7-9)
    ti = data.get("TI_Inspection", [])
    if ti:
        r = ti[0]
        set_cell(table, 9, 2, _num(r.get("TotalPlanned"), 0))
        set_cell(table, 9, 4, _num(r.get("InspectedPrevQtr"), 0))
        set_cell(table, 9, 6, _num(r.get("PlannedThisQtr"), 0))
        set_cell(table, 9, 8, _num(r.get("InspectedThisQtr"), 0))
        set_cell(table, 9, 10, _num(r.get("ArrearIf"), 0))

    # Item 19: IR/Para settlement (Rows 10-12)
    if ti:
        r = ti[0]
        # IRs
        set_cell(table, 12, 2, _num(r.get("OB_IR"), 0))
        set_cell(table, 12, 4, _num(r.get("Add_IR"), 0))
        set_cell(table, 12, 6, _num(r.get("Settled_IR"), 0))
        set_cell(table, 12, 8, _num(r.get("CB_IR"), 0))
        # Paras
        set_cell(table, 12, 3, _num(r.get("OB_Para"), 0))
        set_cell(table, 12, 5, _num(r.get("Add_Para"), 0))
        set_cell(table, 12, 7, _num(r.get("Settled_Para"), 0))
        set_cell(table, 12, 9, _num(r.get("CB_Para"), 0))

    # Item 20: Budget Review (Rows 13-16)
    budget = data.get("Budget_Review", [])
    if budget:
        r = budget[0]
        set_cell(table, 14, 3, _safe(r.get("DatePassed")))
        set_cell(table, 15, 3, _safe(r.get("DateReceived")))
        set_cell(table, 16, 3, _safe(r.get("DateReviewCompleted")))


# ---------------------------------------------------------------------------
# Table 4: GPF Items 1-3 (FP, RB, Missing Credits, Unposted, PF Suspense)
# ---------------------------------------------------------------------------
def populate_table_4(table, data):
    """Fill Table 4 — GPF FP cases, RB cases, Missing/Unposted Credits, PF Suspense."""
    gpf = data.get("GPF_Summary", [])

    # Map KRA_Item to row index in Table 4
    item_row_map = {
        "Final Payments": 4,
        "Residual Balances": 6,
        "Missing Credits": 8,
        "Unposted Credits": 9,
    }

    for rec in gpf:
        kra_item = _safe(rec.get("KRA_Item", ""))
        target_row = None
        for pattern, row_idx in item_row_map.items():
            if pattern.lower() in kra_item.lower():
                target_row = row_idx
                break
        if target_row is None:
            continue

        set_cell(table, target_row, 2, _num(rec.get("OB")))
        set_cell(table, target_row, 3, _num(rec.get("Addition")))
        set_cell(table, target_row, 4, _num(rec.get("Total")))
        clearance = _num(rec.get("Clearance"))
        cpct = rec.get("ClearPct")
        if cpct:
            clearance += f" ({_pct(cpct)})"
        set_cell(table, target_row, 5, clearance)
        set_cell(table, target_row, 6, _num(rec.get("CB")))

    # PF Suspense Credit/Debit (Rows 12-15)
    pf_susp = data.get("GPF_Suspense", [])
    if pf_susp:
        r = pf_susp[0]
        # Credit (Row 13)
        set_cell(table, 13, 2, _num(r.get("CreditAmt")))
        # Debit (Row 15)
        set_cell(table, 15, 2, _num(r.get("DebitAmt")))


# ---------------------------------------------------------------------------
# Table 5: GPF Items 4-9 (Complaints, Posting, Minus, Review, SMS)
# ---------------------------------------------------------------------------
def populate_table_5(table, data):
    """Fill Table 5 — GPF complaints, posting arrears, minus balances, review, SMS."""
    gpf = data.get("GPF_Summary", [])

    # Item 4: Complaint Cases (Rows 0-1)
    for rec in gpf:
        if "complaint" in _safe(rec.get("KRA_Item", "")).lower():
            set_cell(table, 1, 2, _num(rec.get("OB")))
            set_cell(table, 1, 3, _num(rec.get("Addition")))
            set_cell(table, 1, 4, _num(rec.get("Total")))
            clearance = _num(rec.get("Clearance"))
            cpct = rec.get("ClearPct")
            if cpct:
                clearance += f" ({_pct(cpct)})"
            set_cell(table, 1, 5, clearance)
            set_cell(table, 1, 6, _num(rec.get("CB")))

    # Item 5: Posting Arrears (Row 3-4)
    misc = data.get("GPF_Misc", [])
    for rec in misc:
        field = _safe(rec.get("Field", "")).lower()
        val = _safe(rec.get("Value", ""))
        if "posting" in field and "done" in field:
            set_cell(table, 4, 3, val)
        elif "posting" in field and "arrear" in field:
            set_cell(table, 4, 4, val)

    # Item 6: Minus Balances (Rows 4-5)
    for rec in gpf:
        if "minus" in _safe(rec.get("KRA_Item", "")).lower():
            set_cell(table, 5, 3, _num(rec.get("OB")))
            set_cell(table, 5, 4, _num(rec.get("Addition")))
            set_cell(table, 5, 5, _num(rec.get("Total")))
            clearance = _num(rec.get("Clearance"))
            cpct = rec.get("ClearPct")
            if cpct:
                clearance += f" ({_pct(cpct)})"
            set_cell(table, 5, 6, clearance)
            set_cell(table, 5, 7, _num(rec.get("CB")))

    # Items 7-9: GPF closing, slips, review — from GPF_Misc
    for rec in misc:
        field = _safe(rec.get("Field", "")).lower()
        val = _safe(rec.get("Value", ""))
        if "closing" in field and "date" in field:
            set_cell(table, 6, 2, val)
        elif "slip" in field:
            set_cell(table, 7, 2, val)
        elif "review" in field:
            set_cell(table, 9, 3, val)


# ---------------------------------------------------------------------------
# Table 6: GPF Items 10-11 (SMS Registration, Digitization)
# ---------------------------------------------------------------------------
def populate_table_6(table, data):
    """Fill Table 6 — GPF SMS registration and digitization stats."""
    misc = data.get("GPF_Misc", [])
    for rec in misc:
        field = _safe(rec.get("Field", "")).lower()
        val = _safe(rec.get("Value", ""))
        if "sms" in field and "registered" in field:
            set_cell(table, 2, 4, val)
        elif "digitiz" in field and "%" in field:
            set_cell(table, 5, 3, val)
        elif "digitiz" in field and "done" in field:
            set_cell(table, 5, 2, val)


# ---------------------------------------------------------------------------
# Table 7: Annexure A+B (AC Bills & UCs year-wise)
# ---------------------------------------------------------------------------
def populate_table_7(table, data):
    """Fill Table 7 — Year-wise AC Bills (rows 3-5) and UCs (rows 10-21)."""
    ac = data.get("Accounts_ACBills", [])
    uc = data.get("Accounts_UCs", [])

    # AC Bills (rows 3-5 typically)
    for rec in ac:
        year = _safe(rec.get("Year", rec.get("FY", "")))
        row_idx = _find_row_by_text(table, 0, year, start_row=2, end_row=6)
        if row_idx is not None:
            set_cell(table, row_idx, 1, _num(rec.get("OB")))
            set_cell(table, row_idx, 2, _num(rec.get("Addition")))
            set_cell(table, row_idx, 3, _num(rec.get("Total")))
            set_cell(table, row_idx, 4, _num(rec.get("Clearance")))
            set_cell(table, row_idx, 5, _num(rec.get("CB")))

    # UCs Awaited (rows 10-22)
    for rec in uc:
        year = _safe(rec.get("Year", rec.get("FY", "")))
        row_idx = _find_row_by_text(table, 0, year, start_row=8, end_row=22)
        if row_idx is not None:
            set_cell(table, row_idx, 1, _num(rec.get("OB")))
            set_cell(table, row_idx, 2, _num(rec.get("Addition")))
            set_cell(table, row_idx, 3, _num(rec.get("Total")))
            set_cell(table, row_idx, 4, _num(rec.get("Clearance")))
            set_cell(table, row_idx, 5, _num(rec.get("CB")))


# ---------------------------------------------------------------------------
# Tables 8-13: Annexure C (Year-wise Suspense/Remittance breakdowns)
# ---------------------------------------------------------------------------
def _populate_annexure_c(table, data, head_pattern):
    """Generic populator for Annexure C year-wise tables (Tables 8-13)."""
    annex_c = data.get("Suspense_AnnexC", [])
    filtered = [r for r in annex_c if head_pattern in _safe(r.get("Head", ""))]

    for rec in filtered:
        year = _safe(rec.get("Year", ""))
        dr_cr = _safe(rec.get("Dr_Cr", ""))
        row_idx = _find_row_by_text(table, 0, year, start_row=1)
        if row_idx is None:
            continue
        # Adjust row for Dr vs Cr (Dr comes first, Cr on next row)
        if "cr" in dr_cr.lower():
            row_idx += 1

        set_cell(table, row_idx, 1, dr_cr)
        set_cell(table, row_idx, 2, _num(rec.get("OB")))
        set_cell(table, row_idx, 3, _num(rec.get("Addition")))
        set_cell(table, row_idx, 4, _num(rec.get("Total")))
        set_cell(table, row_idx, 5, _num(rec.get("Clearance")))
        set_cell(table, row_idx, 6, _pct(rec.get("ClearPct")))
        set_cell(table, row_idx, 7, _pct(rec.get("OldCleared")))
        set_cell(table, row_idx, 8, _num(rec.get("CB")))


def populate_table_8(table, data):
    _populate_annexure_c(table, data, "8793")       # Inter State Suspense

def populate_table_9(table, data):
    _populate_annexure_c(table, data, "8658-101")   # PAO Suspense

def populate_table_10(table, data):
    _populate_annexure_c(table, data, "8658-109")   # RBS HQ

def populate_table_11(table, data):
    _populate_annexure_c(table, data, "8658-110")   # RBS CAO

def populate_table_12(table, data):
    _populate_annexure_c(table, data, "8782-103")   # Forest Remittances

def populate_table_13(table, data):
    _populate_annexure_c(table, data, "8782-102")   # PW Remittance


# ---------------------------------------------------------------------------
# Table 14: Annexure D (LTA) + Annexure E (TI) combined
# ---------------------------------------------------------------------------
def populate_table_14(table, data):
    """Fill Table 14 — LTA Unposted/Missing (Annexure D) + TI stats (Annexure E)."""

    # Annexure D: LTA details (Rows 2-6)
    lta_d = data.get("LTA_AnnexD", [])
    for rec in lta_d:
        lta_type = _safe(rec.get("Type", "")).lower()
        if "unposted" in lta_type:
            set_cell(table, 4, 2, _num(rec.get("OB_Items")))
            set_cell(table, 4, 3, _num(rec.get("OB_Amt")))
            set_cell(table, 4, 8, _num(rec.get("CB_Items")))
            set_cell(table, 4, 9, _num(rec.get("CB_Amt")))
        elif "missing" in lta_type:
            set_cell(table, 5, 2, _num(rec.get("OB_Items")))
            set_cell(table, 5, 3, _num(rec.get("OB_Amt")))
            set_cell(table, 5, 8, _num(rec.get("CB_Items")))
            set_cell(table, 5, 9, _num(rec.get("CB_Amt")))

    # Annexure E: TI stats (Rows 17-22)
    ti = data.get("TI_Inspection", [])
    if ti:
        r = ti[0]
        set_cell(table, 18, 2, _num(r.get("TotalPlanned"), 0))
        set_cell(table, 18, 3, _num(r.get("InspectedThisQtr"), 0))

    # Annexure E: Year-wise IR/Para (Rows 26-34)
    ti_yr = data.get("TI_AnnexE_YearWise", [])
    for rec in ti_yr:
        year = _safe(rec.get("AccYear", ""))
        row_idx = _find_row_by_text(table, 1, year, start_row=26, end_row=34)
        if row_idx is None:
            continue
        set_cell(table, row_idx, 2, _num(rec.get("OB_IR")))
        set_cell(table, row_idx, 3, _num(rec.get("OB_Para")))
        set_cell(table, row_idx, 4, _num(rec.get("Cleared_IR")))
        set_cell(table, row_idx, 5, _num(rec.get("Cleared_Para")))
        set_cell(table, row_idx, 6, _num(rec.get("CB_IR")))
        set_cell(table, row_idx, 7, _num(rec.get("CB_Para")))
        set_cell(table, row_idx, 8, _safe(rec.get("ActionTaken")))


# ---------------------------------------------------------------------------
# Table 15: Annexure F (CV voucher validation detail)
# ---------------------------------------------------------------------------
def populate_table_15(table, data):
    """Fill Table 15 — Annexure F voucher validation summary."""
    vouchers = data.get("Voucher_Details", [])
    # Month-wise rows (rows 1-4 typically)
    for i, rec in enumerate(vouchers[:4]):
        row_idx = 1 + i
        set_cell(table, row_idx, 0, _num(rec.get("TotalReceived"), 0))
        set_cell(table, row_idx, 1, _num(rec.get("Validated"), 0))
        set_cell(table, row_idx, 3, _pct(rec.get("PctValidated")))


# ---------------------------------------------------------------------------
# Tables 16-21: GPF Annexures H, I, J, K, L
# ---------------------------------------------------------------------------
def populate_table_16(table, data):
    """Annexure H: FP/RB case details."""
    annex_h = data.get("GPF_AnnexH", [])
    row_idx = 3
    for rec in annex_h:
        if row_idx >= len(table.rows):
            break
        set_cell(table, row_idx, 2, _num(rec.get("OB")))
        set_cell(table, row_idx, 3, _num(rec.get("Addition")))
        set_cell(table, row_idx, 4, _num(rec.get("Total")))
        set_cell(table, row_idx, 5, _num(rec.get("Cleared")))
        set_cell(table, row_idx, 6, _pct(rec.get("ClearPct")))
        row_idx += 1


def _populate_gpf_yearwise(table, data, ws_name, start_row=3):
    """Generic populator for year-wise GPF annexure tables."""
    records = data.get(ws_name, [])
    for rec in records:
        year = _safe(rec.get("Year", ""))
        row_idx = _find_row_by_text(table, 1, year, start_row=start_row)
        if row_idx is None:
            continue
        set_cell(table, row_idx, 2, _num(rec.get("OB", rec.get("Items_OB", ""))))
        set_cell(table, row_idx, 3, _num(rec.get("Addition", rec.get("Items_Add", ""))))
        set_cell(table, row_idx, 4, _num(rec.get("Cleared", rec.get("Items_Cleared", ""))))
        set_cell(table, row_idx, 5, _num(rec.get("CB", rec.get("Items_CB", ""))))


def populate_table_17(table, data):
    _populate_gpf_yearwise(table, data, "GPF_AnnexI", start_row=5)  # Missing Credits

def populate_table_18(table, data):
    """Annexure J: Unposted Credits (Items + Amount columns)."""
    records = data.get("GPF_AnnexJ", [])
    for rec in records:
        year = _safe(rec.get("Year", ""))
        row_idx = _find_row_by_text(table, 1, year, start_row=8)
        if row_idx is None:
            continue
        set_cell(table, row_idx, 2, _num(rec.get("Items_OB")))
        set_cell(table, row_idx, 3, _num(rec.get("Amt_OB")))
        set_cell(table, row_idx, 4, _num(rec.get("Items_Cleared")))
        set_cell(table, row_idx, 5, _num(rec.get("Amt_Cleared")))
        set_cell(table, row_idx, 6, _num(rec.get("Items_CB")))
        set_cell(table, row_idx, 7, _num(rec.get("Amt_CB")))


def populate_table_19(table, data):
    """Unposted Debits summary."""
    records = data.get("GPF_AnnexJ", [])
    debit_recs = [r for r in records if "debit" in _safe(r.get("Type", "")).lower()]
    if debit_recs:
        rec = debit_recs[0]
        set_cell(table, 4, 1, _num(rec.get("Items_OB"), "Nil"))
        set_cell(table, 4, 2, _num(rec.get("Amt_OB"), "NA"))


def populate_table_20(table, data):
    """Annexure K: Missing Debits."""
    records = data.get("GPF_AnnexKL", [])
    filtered = [r for r in records if _safe(r.get("Annexure", "")).upper() == "K"]
    for rec in filtered:
        year = _safe(rec.get("Year", ""))
        row_idx = _find_row_by_text(table, 1, year, start_row=5)
        if row_idx is not None:
            set_cell(table, row_idx, 2, _num(rec.get("OB")))
            set_cell(table, row_idx, 3, _num(rec.get("Addition")))
            set_cell(table, row_idx, 4, _num(rec.get("Cleared")))
            set_cell(table, row_idx, 5, _num(rec.get("CB")))


def populate_table_21(table, data):
    """Annexure L: Minus Balances."""
    records = data.get("GPF_AnnexKL", [])
    filtered = [r for r in records if _safe(r.get("Annexure", "")).upper() == "L"]
    for rec in filtered:
        year = _safe(rec.get("Year", ""))
        row_idx = _find_row_by_text(table, 1, year, start_row=7)
        if row_idx is not None:
            set_cell(table, row_idx, 2, _num(rec.get("OB")))
            set_cell(table, row_idx, 3, _num(rec.get("Cleared")))
            set_cell(table, row_idx, 4, _num(rec.get("CB")))


# ---------------------------------------------------------------------------
# Table 22: DPF Details
# ---------------------------------------------------------------------------
def populate_table_22(table, data):
    """Fill DPF details."""
    dpf = data.get("GPF_DPF", [])
    if dpf:
        r = dpf[0]
        set_cell(table, 2, 0, _num(r.get("OB")))
        set_cell(table, 2, 1, _num(r.get("Addition")))
        set_cell(table, 2, 2, _num(r.get("Total")))
        set_cell(table, 2, 3, _num(r.get("Clearance")))
        set_cell(table, 2, 4, _num(r.get("CB")))


# ---------------------------------------------------------------------------
# Table 23: CV category-wise voucher detail
# ---------------------------------------------------------------------------
def populate_table_23(table, data):
    """Fill Annexure F category-wise voucher validation."""
    vouchers = data.get("Voucher_Details", [])
    type_row_map = {
        "pay bill": 2, "fvc": 3, "gia": 4, "ta": 5,
        "medical": 6, "pension": 7, "refund": 8, "others": 9, "misc": 9
    }
    for rec in vouchers:
        v_type = _safe(rec.get("VoucherType", "")).lower()
        row_idx = None
        for pattern, ridx in type_row_map.items():
            if pattern in v_type:
                row_idx = ridx
                break
        if row_idx is None:
            continue
        set_cell(table, row_idx, 2, _num(rec.get("Validated"), 0))
        set_cell(table, row_idx, 3, _num(rec.get("TotalReceived"), 0))
        set_cell(table, row_idx, 4, _pct(rec.get("PctValidated")))


# ---------------------------------------------------------------------------
# Tables 24-31: Admin Complaints, RTI, Court Cases (AG-I, AG-II, Bhopal)
# ---------------------------------------------------------------------------
def _populate_complaints_table(table, data, office_filter):
    """Populate a complaints/RTI table for a specific office."""
    complaints = _get_records(data, "Complaints_RTI", Office=office_filter)
    row_map = {"direct": 1, "cag": 2, "rti": 3, "court": 4}
    for rec in complaints:
        comp_type = _safe(rec.get("Type", "")).lower()
        row_idx = None
        for pattern, ridx in row_map.items():
            if pattern in comp_type:
                row_idx = ridx
                break
        if row_idx is None:
            continue
        set_cell(table, row_idx, 1, _num(rec.get("OB")))
        set_cell(table, row_idx, 2, _num(rec.get("Received")))
        set_cell(table, row_idx, 3, _num(rec.get("DisposedInTime")))
        set_cell(table, row_idx, 4, _num(rec.get("DisposedBeyond")))
        set_cell(table, row_idx, 5, _num(rec.get("CB")))
        set_cell(table, row_idx, 6, _safe(rec.get("DelayReason")))
        set_cell(table, row_idx, 7, _safe(rec.get("Categories")))


def _populate_court_age_table(table, data, office_filter):
    """Populate age-wise court cases table."""
    age_data = _get_records(data, "Court_AgeWise", Office=office_filter)
    age_row_map = {"10": 1, "5-10": 2, "2-5": 3, "less": 4, "<2": 4}
    for rec in age_data:
        age = _safe(rec.get("AgeGroup", "")).lower()
        row_idx = None
        for pattern, ridx in age_row_map.items():
            if pattern in age:
                row_idx = ridx
                break
        if row_idx is not None:
            set_cell(table, row_idx, 1, _num(rec.get("CaseCount")))
            set_cell(table, row_idx, 2, _safe(rec.get("Remarks")))


def populate_table_24(table, data):
    _populate_complaints_table(table, data, "AG-I")

def populate_table_25(table, data):
    _populate_court_age_table(table, data, "AG-I")

def populate_table_26(table, data):
    """AG-II complaints."""
    complaints = _get_records(data, "Complaints_RTI", Office="AG-II")
    for i, rec in enumerate(complaints):
        row_idx = 1 + i
        if row_idx >= len(table.rows):
            break
        set_cell(table, row_idx, 1, _num(rec.get("OB")))
        set_cell(table, row_idx, 2, _num(rec.get("Received")))
        set_cell(table, row_idx, 3, _num(rec.get("DisposedInTime")))
        set_cell(table, row_idx, 4, _num(rec.get("DisposedBeyond")))
        set_cell(table, row_idx, 5, _num(rec.get("CB")))

def populate_table_27(table, data):
    """AG-II RTI cases."""
    rti = _get_records(data, "Complaints_RTI", Office="AG-II")
    rti = [r for r in rti if "rti" in _safe(r.get("Type", "")).lower()]
    if rti:
        rec = rti[0]
        set_cell(table, 1, 1, _num(rec.get("OB")))
        set_cell(table, 1, 2, _num(rec.get("Received")))
        set_cell(table, 1, 3, _num(rec.get("DisposedInTime")))
        set_cell(table, 1, 4, _num(rec.get("DisposedBeyond")))
        set_cell(table, 1, 5, _num(rec.get("CB")))

def populate_table_28(table, data):
    """AG-II Court Cases."""
    courts = _get_records(data, "Court_Cases", Office="AG-II")
    type_row = {"supreme": 1, "high": 2, "cat": 3, "lower": 4, "subordinate": 4}
    for rec in courts:
        ct = _safe(rec.get("CourtType", "")).lower()
        row_idx = None
        for pattern, ridx in type_row.items():
            if pattern in ct:
                row_idx = ridx
                break
        if row_idx is not None:
            set_cell(table, row_idx, 1, _num(rec.get("OB")))
            set_cell(table, row_idx, 2, _num(rec.get("Added")))
            set_cell(table, row_idx, 3, _num(rec.get("Decided")))
            set_cell(table, row_idx, 4, _num(rec.get("CB")))

def populate_table_29(table, data):
    _populate_court_age_table(table, data, "AG-II")

def populate_table_30(table, data):
    _populate_complaints_table(table, data, "Bhopal")

def populate_table_31(table, data):
    _populate_court_age_table(table, data, "Bhopal")


# ---------------------------------------------------------------------------
# Tables 32-34: VLC Annexure U (Automation, Disruptions, Change Mgmt)
# ---------------------------------------------------------------------------
def populate_table_32(table, data):
    """Annexure U: VLC automation listing."""
    auto = data.get("VLC_Automation", [])
    for i, rec in enumerate(auto):
        row_idx = 1 + i
        if row_idx >= len(table.rows):
            break
        set_cell(table, row_idx, 1, _safe(rec.get("AppName")))
        set_cell(table, row_idx, 2, _safe(rec.get("Database")))
        set_cell(table, row_idx, 3, _safe(rec.get("OS")))
        set_cell(table, row_idx, 4, _safe(rec.get("AppPlatform")))
        set_cell(table, row_idx, 5, _safe(rec.get("QueryPlatform")))
        set_cell(table, row_idx, 6, _safe(rec.get("DataSource")))
        set_cell(table, row_idx, 7, _safe(rec.get("DataInput")))
        set_cell(table, row_idx, 8, _safe(rec.get("OutputForm")))
        set_cell(table, row_idx, 9, _safe(rec.get("Backup")))


def populate_table_33(table, data):
    """VLC Service Disruptions."""
    disruptions = data.get("VLC_ServiceDisruption", [])
    for i, rec in enumerate(disruptions):
        row_idx = 2 + i
        if row_idx >= len(table.rows):
            break
        set_cell(table, row_idx, 1, _safe(rec.get("DaysDisrupted")))
        set_cell(table, row_idx, 2, _safe(rec.get("Reason")))
        set_cell(table, row_idx, 3, _safe(rec.get("RestorationDate")))
        set_cell(table, row_idx, 4, _safe(rec.get("ActionTaken")))


def populate_table_34(table, data):
    """VLC Change Management."""
    changes = data.get("VLC_ChangeMgmt", [])
    for i, rec in enumerate(changes):
        row_idx = 1 + i
        if row_idx >= len(table.rows):
            break
        set_cell(table, row_idx, 0, _safe(rec.get("Function")))
        set_cell(table, row_idx, 1, _safe(rec.get("Description")))
        set_cell(table, row_idx, 2, _safe(rec.get("CompletionDate")))
        set_cell(table, row_idx, 3, _safe(rec.get("PendingProposal")))
        set_cell(table, row_idx, 4, _safe(rec.get("Remarks")))


# ---------------------------------------------------------------------------
# Table 35: Annexure W (TI digitization)
# ---------------------------------------------------------------------------
def populate_table_35(table, data):
    """TI Cell Treasury digitization questions."""
    misc = data.get("GPF_Misc", [])  # Using misc for now; can be TI-specific
    # These are typically static text answers, populated row-by-row
    # Rows 0-3: 4 digitization questions
    pass  # Populated manually or via a dedicated TI_Digitization worksheet


# ---------------------------------------------------------------------------
# Table 36: Annexure X (Broadsheet differences)
# ---------------------------------------------------------------------------
def populate_table_36(table, data):
    """Fill Broadsheet vs Ledger differences."""
    diffs = data.get("Broadsheet_Diff", [])
    head_row_map = {
        "loan": 1, "8443": 2, "deposit": 3,
        "8658": 4, "suspense": 4,
        "8782-102": 5, "remittance": 5,
        "forest": 5, "8782-103": 5
    }
    for rec in diffs:
        head = _safe(rec.get("HeadOfAccount", "")).lower()
        row_idx = None
        for pattern, ridx in head_row_map.items():
            if pattern in head:
                row_idx = ridx
                break
        if row_idx is not None:
            set_cell(table, row_idx, 1, _num(rec.get("DiffAmount")))
            set_cell(table, row_idx, 2, _num(rec.get("ClearanceDuringQtr")))


# ---------------------------------------------------------------------------
# Tables 37-39: GPF stats (Incomplete accounts, PF Suspense, Online receipt)
# ---------------------------------------------------------------------------
def populate_table_37(table, data):
    """Incomplete PF accounts stats."""
    misc = data.get("GPF_Misc", [])
    for rec in misc:
        field = _safe(rec.get("Field", "")).lower()
        val = _safe(rec.get("Value", ""))
        if "live" in field or "total" in field:
            set_cell(table, 1, 0, val)
        elif "complete" in field and "in" not in field:
            set_cell(table, 1, 1, val)
        elif "incomplete" in field:
            set_cell(table, 1, 2, val)
        elif "dormant" in field:
            set_cell(table, 1, 3, val)


def populate_table_38(table, data):
    """PF Suspense clearance amounts."""
    pf_susp = data.get("GPF_Suspense", [])
    if pf_susp:
        r = pf_susp[0]
        set_cell(table, 1, 1, _num(r.get("CreditAmt")))
        set_cell(table, 1, 2, _num(r.get("DebitAmt")))


def populate_table_39(table, data):
    """Online GPF Credit/Debit receipt."""
    online = data.get("GPF_Online", [])
    for rec in online:
        gtype = _safe(rec.get("Type", "")).lower()
        month = _safe(rec.get("Month", ""))
        if "credit" in gtype:
            row_idx = _find_row_by_text(table, 0, "A)", start_row=0)
            if row_idx is not None:
                set_cell(table, row_idx + 1, 2, _num(rec.get("Subscribers")))
                set_cell(table, row_idx + 1, 3, _num(rec.get("EntriesDue")))
                set_cell(table, row_idx + 1, 4, _num(rec.get("ReceivedOnline")))
                set_cell(table, row_idx + 1, 5, _pct(rec.get("PctOnline")))
        elif "debit" in gtype:
            row_idx = _find_row_by_text(table, 0, "B)", start_row=0)
            if row_idx is not None:
                set_cell(table, row_idx + 1, 2, _num(rec.get("Subscribers")))
                set_cell(table, row_idx + 1, 3, _num(rec.get("EntriesDue")))
                set_cell(table, row_idx + 1, 4, _num(rec.get("ReceivedOnline")))
                set_cell(table, row_idx + 1, 5, _pct(rec.get("PctOnline")))


# ---------------------------------------------------------------------------
# Tables 40-41: Deposit PDA details
# ---------------------------------------------------------------------------
def populate_table_40(table, data):
    """PDA summary (9 columns)."""
    pda = data.get("Deposit_PDA", [])
    if pda:
        r = pda[0]
        set_cell(table, 1, 0, _num(r.get("TotalPDAs"), 0))
        set_cell(table, 1, 1, _num(r.get("BalAmt")))
        set_cell(table, 1, 2, _num(r.get("Closed"), 0))
        set_cell(table, 1, 3, _num(r.get("ClosedAmt")))
        set_cell(table, 1, 4, _num(r.get("Opened"), 0))
        set_cell(table, 1, 5, _num(r.get("OpenedAmt")))
        set_cell(table, 1, 6, _num(r.get("PermissionSought"), 0))
        set_cell(table, 1, 7, _num(r.get("Inoperative"), 0))
        set_cell(table, 1, 8, _num(r.get("InoperativeAmt")))


def populate_table_41(table, data):
    """Closed PD Account details — dynamic rows."""
    # This table has variable rows, typically populated from separate data
    pass  # Populated from a future Deposit_Closed worksheet if created


# ---------------------------------------------------------------------------
# Table 42: Annexure Y (ARU Statistics)
# ---------------------------------------------------------------------------
def populate_table_42(table, data):
    """ARU statistical information (13 categories)."""
    aru = data.get("VLC_ARU", [])
    cat_row_map = {
        "treasuries": 1, "sub-treasuries": 2, "voucher": 3,
        "challan": 4, "pw": 5, "public health": 6, "forest": 7,
        "irrigation": 8, "res": 9, "narmada": 9, "cco": 10,
        "bco": 10, "ddo": 11, "non-banking": 12
    }
    for rec in aru:
        cat = _safe(rec.get("Category", "")).lower()
        row_idx = None
        for pattern, ridx in cat_row_map.items():
            if pattern in cat:
                row_idx = ridx
                break
        if row_idx is not None:
            set_cell(table, row_idx, 1, _num(rec.get("Numbers"), 0))


# ---------------------------------------------------------------------------
# Tables 43-44: Staff Strength
# ---------------------------------------------------------------------------
def populate_table_43(table, data):
    """SS vs PIP cadre-wise."""
    staff = data.get("Staff_Strength", [])
    for rec in staff:
        cadre = _safe(rec.get("Cadre", ""))
        row_idx = _find_row_by_text(table, 0, cadre, start_row=2)
        if row_idx is not None:
            set_cell(table, row_idx, 1, _num(rec.get("SS"), 0))
            set_cell(table, row_idx, 2, _num(rec.get("PIP"), 0))
            set_cell(table, row_idx, 3, _pct(rec.get("VacancyPct")))


def populate_table_44(table, data):
    """Group-wise distribution."""
    group = data.get("Staff_GroupWise", [])
    for rec in group:
        cadre = _safe(rec.get("Cadre", ""))
        row_idx = _find_row_by_text(table, 0, cadre, start_row=3)
        if row_idx is not None:
            set_cell(table, row_idx, 1, _num(rec.get("Permanent"), 0))
            set_cell(table, row_idx, 2, _num(rec.get("Temporary"), 0))
            set_cell(table, row_idx, 3, _num(rec.get("Casual"), 0))
            set_cell(table, row_idx, 4, _num(rec.get("AccountsGrp"), 0))
            set_cell(table, row_idx, 5, _num(rec.get("FundsGrp"), 0))
            set_cell(table, row_idx, 6, _num(rec.get("BhopalBranch"), 0))
            set_cell(table, row_idx, 7, _num(rec.get("AdminGrp"), 0))
            set_cell(table, row_idx, 8, _num(rec.get("OthersGE"), 0))
            set_cell(table, row_idx, 9, _num(rec.get("Total"), 0))


# ---------------------------------------------------------------------------
# Tables 45-46: IFMS Status (Annexure Z)
# ---------------------------------------------------------------------------
def populate_table_45(table, data):
    """IFMS module status (first batch)."""
    ifms = data.get("IFMS_Status", [])
    for i, rec in enumerate(ifms[:6]):
        row_idx = 1 + i
        if row_idx >= len(table.rows):
            break
        set_cell(table, row_idx, 2, _safe(rec.get("ImplementationStatus")))
        set_cell(table, row_idx, 3, _safe(rec.get("OnlineData")))
        set_cell(table, row_idx, 4, _safe(rec.get("InterfaceDeveloped")))


def populate_table_46(table, data):
    """IFMS module status (second batch)."""
    ifms = data.get("IFMS_Status", [])
    for i, rec in enumerate(ifms[6:]):
        row_idx = i
        if row_idx >= len(table.rows):
            break
        set_cell(table, row_idx, 2, _safe(rec.get("ImplementationStatus")))
        set_cell(table, row_idx, 3, _safe(rec.get("OnlineData")))
        set_cell(table, row_idx, 4, _safe(rec.get("InterfaceDeveloped")))


# ---------------------------------------------------------------------------
# Table 47: IFMS online ARU receipt
# ---------------------------------------------------------------------------
def populate_table_47(table, data):
    """IFMS online account receipt from ARUs."""
    ifms_online = data.get("IFMS_Online", [])
    for i, rec in enumerate(ifms_online):
        row_idx = 1 + i
        if row_idx >= len(table.rows):
            break
        set_cell(table, row_idx, 2, _num(rec.get("ARU_Total")))
        set_cell(table, row_idx, 3, _num(rec.get("ARU_Online")))
        set_cell(table, row_idx, 4, _pct(rec.get("PctOnline")))
        set_cell(table, row_idx, 5, _num(rec.get("ARU_Pending")))
        set_cell(table, row_idx, 6, _safe(rec.get("Remarks")))


# ===========================================================================
# Master function: populate ALL tables
# ===========================================================================

# Registry of all table populators (index -> function)
TABLE_POPULATORS = {
    0: populate_table_0,
    1: populate_table_1,
    2: populate_table_2,
    3: populate_table_3,
    4: populate_table_4,
    5: populate_table_5,
    6: populate_table_6,
    7: populate_table_7,
    8: populate_table_8,
    9: populate_table_9,
    10: populate_table_10,
    11: populate_table_11,
    12: populate_table_12,
    13: populate_table_13,
    14: populate_table_14,
    15: populate_table_15,
    16: populate_table_16,
    17: populate_table_17,
    18: populate_table_18,
    19: populate_table_19,
    20: populate_table_20,
    21: populate_table_21,
    22: populate_table_22,
    23: populate_table_23,
    24: populate_table_24,
    25: populate_table_25,
    26: populate_table_26,
    27: populate_table_27,
    28: populate_table_28,
    29: populate_table_29,
    30: populate_table_30,
    31: populate_table_31,
    32: populate_table_32,
    33: populate_table_33,
    34: populate_table_34,
    35: populate_table_35,
    36: populate_table_36,
    37: populate_table_37,
    38: populate_table_38,
    39: populate_table_39,
    40: populate_table_40,
    41: populate_table_41,
    42: populate_table_42,
    43: populate_table_43,
    44: populate_table_44,
    45: populate_table_45,
    46: populate_table_46,
    47: populate_table_47,
}


def populate_all_tables(doc, data):
    """
    Populate all tables in the KRA Word document with data.

    Args:
        doc: python-docx Document object (the blank template)
        data: dict from fetch_data.fetch_all_data() — {worksheet_name: [records]}
    """
    total_tables = len(doc.tables)
    print(f"Template has {total_tables} tables. Populating...")

    populated = 0
    skipped = 0

    for idx in range(total_tables):
        populator = TABLE_POPULATORS.get(idx)
        if populator:
            try:
                populator(doc.tables[idx], data)
                populated += 1
                print(f"  Table {idx:2d}: OK ({populator.__name__})")
            except Exception as e:
                print(f"  Table {idx:2d}: ERROR in {populator.__name__} - {e}")
                skipped += 1
        else:
            print(f"  Table {idx:2d}: SKIPPED (no populator defined)")
            skipped += 1

    print(f"\nPopulation complete: {populated} tables filled, {skipped} skipped/errored")
