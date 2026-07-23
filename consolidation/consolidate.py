"""
consolidate.py -- KRA Report Consolidation Engine (Main Orchestrator)
============================================================================
Reads all section-wise KRA data from Google Sheets, fills the blank
KRA Word template, generates the Appreciation Note, and saves the
compiled report as a .docx file.

Usage:
    py consolidate.py --fy 2026-27 --quarter Q1
    py consolidate.py --fy 2026-27 --quarter Q1 --template "path/to/template.docx"
    py consolidate.py --fy 2026-27 --quarter Q1 --dry-run

Output:
    output/Consolidated_KRA_Q1_2026-27.docx
"""

import os
import sys
import argparse
from docx import Document

from fetch_data import fetch_all_data
from populate_template import populate_all_tables
from appreciation_note import generate_appreciation_note, insert_appreciation_note


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
DEFAULT_TEMPLATE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "Sectionwise Blank KRA Report for June 2026 end Quarter.docx"
)
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")

# Quarter display names
QUARTER_NAMES = {
    "Q1": "June Ending",
    "Q2": "September Ending",
    "Q3": "December Ending",
    "Q4": "March Ending",
}


def main():
    parser = argparse.ArgumentParser(
        description="Compile Section-wise KRA data into the Master KRA Report",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  py consolidate.py --fy 2026-27 --quarter Q1
  py consolidate.py --fy 2026-27 --quarter Q1 --dry-run
  py consolidate.py --fy 2026-27 --quarter Q1 --template "custom_template.docx"
        """
    )
    parser.add_argument("--fy", required=True, help="Financial Year (e.g. 2026-27)")
    parser.add_argument("--quarter", required=True, choices=["Q1", "Q2", "Q3", "Q4"],
                        help="Quarter (Q1, Q2, Q3, or Q4)")
    parser.add_argument("--template", default=DEFAULT_TEMPLATE,
                        help="Path to the blank KRA Word template")
    parser.add_argument("--creds", default=None,
                        help="Path to Google service account credentials.json")
    parser.add_argument("--output", default=None,
                        help="Output file path (default: output/Consolidated_KRA_<Q>_<FY>.docx)")
    parser.add_argument("--dry-run", action="store_true",
                        help="Test Google Sheets connection only, don't generate report")
    parser.add_argument("--skip-note", action="store_true",
                        help="Skip generating the Appreciation Note")
    args = parser.parse_args()

    qtr_name = QUARTER_NAMES.get(args.quarter, args.quarter)
    print("=" * 70)
    print(f"  KRA CONSOLIDATION ENGINE")
    print(f"  Financial Year: {args.fy}")
    print(f"  Quarter: {args.quarter} ({qtr_name})")
    print("=" * 70)
    print()

    # -----------------------------------------------------------------------
    # Step 1: Fetch data from Google Sheets
    # -----------------------------------------------------------------------
    print("[STEP 1] Fetching data from Google Sheets...")
    data = fetch_all_data(args.fy, args.quarter, creds_path=args.creds, dry_run=args.dry_run)

    if args.dry_run:
        print("\nDry run complete. No report generated.")
        return

    # Check if we have any data at all
    total_rows = sum(len(v) for v in data.values())
    if total_rows == 0:
        print("\nWARNING: No data found for the specified FY and Quarter.")
        print("The report will be generated with empty fields.")
        response = input("Continue? (y/n): ").strip().lower()
        if response != "y":
            print("Aborted.")
            return

    # -----------------------------------------------------------------------
    # Step 2: Open Word template
    # -----------------------------------------------------------------------
    print(f"\n[STEP 2] Opening Word template...")
    if not os.path.exists(args.template):
        print(f"ERROR: Template not found at: {args.template}")
        print("Provide the correct path with --template")
        sys.exit(1)

    doc = Document(args.template)
    print(f"  Template loaded: {os.path.basename(args.template)}")
    print(f"  Tables found: {len(doc.tables)}")
    print(f"  Paragraphs found: {len(doc.paragraphs)}")

    # -----------------------------------------------------------------------
    # Step 3: Populate all tables
    # -----------------------------------------------------------------------
    print(f"\n[STEP 3] Populating tables with data...")
    populate_all_tables(doc, data)

    # -----------------------------------------------------------------------
    # Step 4: Generate Appreciation Note
    # -----------------------------------------------------------------------
    if not args.skip_note:
        print(f"\n[STEP 4] Generating Appreciation Note...")
        note_paragraphs = generate_appreciation_note(data, args.fy, args.quarter)
        if note_paragraphs:
            insert_appreciation_note(doc, note_paragraphs)
            print(f"  Generated {len(note_paragraphs)} paragraphs")
        else:
            print("  No data available for appreciation note")
    else:
        print(f"\n[STEP 4] Skipping Appreciation Note (--skip-note)")

    # -----------------------------------------------------------------------
    # Step 5: Save compiled report
    # -----------------------------------------------------------------------
    print(f"\n[STEP 5] Saving compiled report...")

    if args.output:
        out_path = args.output
    else:
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        out_filename = f"Consolidated_KRA_{args.quarter}_{args.fy.replace('/', '-')}.docx"
        out_path = os.path.join(OUTPUT_DIR, out_filename)

    doc.save(out_path)
    file_size_kb = os.path.getsize(out_path) / 1024

    print()
    print("=" * 70)
    print(f"  COMPILATION COMPLETE!")
    print(f"  Output: {os.path.abspath(out_path)}")
    print(f"  Size: {file_size_kb:.1f} KB")
    print(f"  Tables populated: {len(doc.tables)}")
    print("=" * 70)
    print()
    print("Next steps:")
    print("  1. Open the report in Microsoft Word")
    print("  2. Review all tables for accuracy")
    print("  3. Edit the Appreciation Note as needed")
    print("  4. Save as PDF for final distribution")


def run_compilation(fy, quarter, template_path=None, output_path=None, skip_note=False):
    """
    Programmatic entry point for KRA consolidation.
    Returns (success_bool, message_str)
    """
    try:
        q_name = quarter.split(' ')[0] # e.g. "Q1" from "Q1 (June Ending)"
        t_path = template_path or DEFAULT_TEMPLATE
        
        # 1. Fetch data
        print(f"Server-triggered fetch for {q_name} FY {fy}...")
        data = fetch_all_data(fy, q_name)
        
        # 2. Open template
        if not os.path.exists(t_path):
            return False, f"Template not found at: {t_path}"
        doc = Document(t_path)
        
        # 3. Populate
        populate_all_tables(doc, data)
        
        # 4. Appreciation note
        if not skip_note:
            note_paragraphs = generate_appreciation_note(data, fy, q_name)
            if note_paragraphs:
                insert_appreciation_note(doc, note_paragraphs)
                
        # 5. Save output
        if output_path:
            out_path = output_path
        else:
            os.makedirs(OUTPUT_DIR, exist_ok=True)
            out_filename = f"Consolidated_KRA_{q_name}_{fy.replace('/', '-')}.docx"
            out_path = os.path.join(OUTPUT_DIR, out_filename)
            
        doc.save(out_path)
        return True, f"Consolidated KRA compiled successfully! Saved to: {os.path.basename(out_path)}"
        
    except Exception as e:
        import traceback
        err_msg = f"Compilation failed: {e}\n{traceback.format_exc()}"
        print(err_msg)
        return False, err_msg


if __name__ == "__main__":
    main()
